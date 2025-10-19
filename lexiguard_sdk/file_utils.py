# lexiguard_sdk/file_utils.py
"""
File parsing utilities for LexiGuard SDK
Supports PDF, DOCX, and TXT files
"""

from typing import Dict, Any, Union
from pathlib import Path
import io


class FileParsingError(Exception):
    """Exception raised when file parsing fails"""
    pass


class FileParser:
    """
    Utility class for parsing different file formats.
    """
    
    @staticmethod
    def parse_pdf(file_path: Union[str, Path, io.BytesIO]) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file or BytesIO object
            
        Returns:
            Extracted text content
            
        Raises:
            FileParsingError: If PDF parsing fails
        """
        try:
            import PyPDF2
        except ImportError:
            raise FileParsingError(
                "PyPDF2 is not installed. Install with: pip install PyPDF2"
            )
        
        try:
            if isinstance(file_path, io.BytesIO):
                reader = PyPDF2.PdfReader(file_path)
            else:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text_content = []
                    for page in reader.pages:
                        text_content.append(page.extract_text())
                    return "\n".join(text_content)
            
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())
            return "\n".join(text_content)
            
        except Exception as e:
            raise FileParsingError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: Union[str, Path, io.BytesIO]) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file or BytesIO object
            
        Returns:
            Extracted text content
            
        Raises:
            FileParsingError: If DOCX parsing fails
        """
        try:
            import docx
        except ImportError:
            raise FileParsingError(
                "python-docx is not installed. Install with: pip install python-docx"
            )
        
        try:
            if isinstance(file_path, io.BytesIO):
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise FileParsingError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(file_path: Union[str, Path, io.BytesIO]) -> str:
        """
        Read text from TXT file.
        
        Args:
            file_path: Path to TXT file or BytesIO object
            
        Returns:
            File content
            
        Raises:
            FileParsingError: If TXT reading fails
        """
        try:
            if isinstance(file_path, io.BytesIO):
                return file_path.read().decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
                    
        except Exception as e:
            raise FileParsingError(f"Failed to read TXT: {str(e)}")
    
    @staticmethod
    def parse_file(file_path: Union[str, Path], file_type: str = None) -> Dict[str, Any]:
        """
        Auto-detect and parse file based on extension.
        
        Args:
            file_path: Path to file
            file_type: Optional file type override ('pdf', 'docx', 'txt')
            
        Returns:
            Dictionary with success status and extracted text
        """
        path = Path(file_path)
        
        # Determine file type
        if file_type is None:
            extension = path.suffix.lower()
            if extension == '.pdf':
                file_type = 'pdf'
            elif extension in ['.docx', '.doc']:
                file_type = 'docx'
            elif extension == '.txt':
                file_type = 'txt'
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {extension}"
                }
        
        # Parse based on type
        try:
            if file_type == 'pdf':
                text = FileParser.parse_pdf(file_path)
            elif file_type == 'docx':
                text = FileParser.parse_docx(file_path)
            elif file_type == 'txt':
                text = FileParser.parse_txt(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_type}"
                }
            
            return {
                "success": True,
                "text": text,
                "file_type": file_type,
                "file_name": path.name
            }
            
        except FileParsingError as e:
            return {
                "success": False,
                "error": str(e)
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Unexpected error: {str(e)}"
            }
    
    @staticmethod
    def parse_uploaded_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse file from uploaded bytes (useful for web frameworks).
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to detect type)
            
        Returns:
            Dictionary with success status and extracted text
        """
        file_obj = io.BytesIO(file_bytes)
        
        # Detect file type from filename
        extension = Path(filename).suffix.lower()
        
        try:
            if extension == '.pdf':
                text = FileParser.parse_pdf(file_obj)
                file_type = 'pdf'
            elif extension in ['.docx', '.doc']:
                text = FileParser.parse_docx(file_obj)
                file_type = 'docx'
            elif extension == '.txt':
                text = FileParser.parse_txt(file_obj)
                file_type = 'txt'
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {extension}"
                }
            
            return {
                "success": True,
                "text": text,
                "file_type": file_type,
                "file_name": filename
            }
            
        except FileParsingError as e:
            return {
                "success": False,
                "error": str(e)
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Unexpected error: {str(e)}"
            }


# Convenience function for quick file analysis
def analyze_file_quick(api_key: str, file_path: str) -> Dict[str, Any]:
    """
    Quick utility to parse and analyze a file in one call.
    
    Args:
        api_key: Google Gemini API key
        file_path: Path to the file
        
    Returns:
        Analysis results
    """
    from .core import LexiGuard
    
    # Parse file
    parse_result = FileParser.parse_file(file_path)
    if not parse_result["success"]:
        return parse_result
    
    # Analyze text
    lg = LexiGuard(api_key=api_key)
    analysis = lg.analyze_text(parse_result["text"])
    
    # Add file info to result
    if analysis["success"]:
        analysis["file_info"] = {
            "file_name": parse_result["file_name"],
            "file_type": parse_result["file_type"]
        }
    
    return analysis