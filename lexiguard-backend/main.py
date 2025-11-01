﻿# main.py - FIXED VERSION

import os
import json
import io
import logging
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict
from typing import Optional, List, Dict
from dotenv import load_dotenv
import google.generativeai as genai
from google.cloud import dlp_v2
from google.cloud.dlp_v2 import types as dlp_types
import PyPDF2
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

from google.cloud import storage as gcs_storage
import uuid
from datetime import datetime

# --- 0. CONFIGURE LOGGING ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- 1. LOAD ENVIRONMENT VARIABLES ---
load_dotenv()

# Set Google Cloud credentials path
credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if credentials_path and not os.path.isabs(credentials_path):
    # Convert relative path to absolute path
    credentials_path = os.path.join(os.path.dirname(__file__), credentials_path)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = credentials_path
    logger.info(f" Google Cloud credentials set to: {credentials_path}")

# --- 2. INITIALIZE FASTAPI APP ---
app = FastAPI(
    title="LexiGuard API",
    description="Analyzes legal documents (text, PDF, or DOCX) using Google's Gemini AI with PII Redaction.",
    version="1.4.0"
)

# --- 3. ENABLE CORS ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://lexiguard-frontend-372716482731.asia-south1.run.app",
        "http://localhost:3000",
        "http://localhost:8000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# --- 4. CONFIGURE GOOGLE GEMINI ---
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    logger.error(" GOOGLE_API_KEY not found in .env file!")
    logger.error("Please add GOOGLE_API_KEY=your_api_key to your .env file")
    raise Exception("GOOGLE_API_KEY is required for Gemini API access")
else:
    # Configure Gemini with explicit API key to avoid credential conflicts
    genai.configure(api_key=API_KEY)
    logger.info(f"Gemini API configured with API key: {API_KEY[:8]}...")

safety_settings = {
    "HARM_CATEGORY_HARASSMENT": "block_none",
    "HARM_CATEGORY_HATE_SPEECH": "block_none", 
    "HARM_CATEGORY_SEXUALLY_EXPLICIT": "block_none",
    "HARM_CATEGORY_DANGEROUS_CONTENT": "block_none",
}

# Initialize Gemini model with proper model discovery
model = None
MODEL_NAME = None

def initialize_gemini_model():
    """Initialize Gemini model with fixed Gemini 2.5 Flash"""
    global model, MODEL_NAME
    
    try:
        # Use fixed Gemini 2.5 Flash model for consistent performance
        MODEL_NAME = "models/gemini-2.5-flash"
        logger.info(f"ï¿½ Initializing fixed model: {MODEL_NAME}")
        
        try:
            # Initialize Gemini 2.5 Flash directly
            test_model = genai.GenerativeModel(MODEL_NAME, safety_settings=safety_settings)
            
            # Test the model with a simple prompt
            test_response = test_model.generate_content("Hello! Please respond with 'Working'.")
            
            if test_response and test_response.text:
                model = test_model
                logger.info(f" Gemini 2.5 Flash initialized successfully: {MODEL_NAME}")
                logger.info(f"Test response: {test_response.text[:50]}...")
                return True
            else:
                raise Exception("Model test failed - no response received")
                
        except Exception as flash_error:
            logger.warning(f" Gemini 2.5 Flash failed: {flash_error}")
            
            # Fallback to models/ prefix format if needed
            try:
                MODEL_NAME = "models/gemini-flash-latest"
                logger.info(f"ðŸ”„ Trying fallback format: {MODEL_NAME}")
                test_model = genai.GenerativeModel(MODEL_NAME, safety_settings=safety_settings)
                test_response = test_model.generate_content("Hello! Please respond with 'Working'.")
                
                if test_response and test_response.text:
                    model = test_model
                    logger.info(f" Gemini 2.5 Flash (fallback format) initialized: {MODEL_NAME}")
                    return True
                    
            except Exception as fallback_error:
                logger.error(f" Both Gemini 2.5 Flash formats failed: {fallback_error}")
                return False
        
        # If we reach here, no model worked
        logger.error(" No working Gemini model found!")
        logger.error("Available models found:", available_models)
        return False
        
    except Exception as e:
        logger.error(f" Failed to initialize any Gemini model: {e}")
        return False

# Initialize the model
success = initialize_gemini_model()
if not success:
    logger.error(" CRITICAL: Gemini model initialization failed completely")
    logger.error("Possible solutions:")
    logger.error("  1. Verify your GOOGLE_API_KEY is correct and has Gemini access")
    logger.error("  2. Check if you have billing enabled for Gemini API")
    logger.error("  3. Ensure the API key has the correct scopes")
    logger.error("  4. Try regenerating your API key in Google Cloud Console")
    logger.error("  5. Verify Generative AI API is enabled in Google Cloud Console")

# --- 5. CONFIGURE GOOGLE CLOUD DLP ---
# Initialize DLP client lazily to avoid multiprocessing issues
dlp_client = None
PROJECT_ID = os.getenv("GOOGLE_CLOUD_PROJECT")

def get_dlp_client():
    """Lazy initialization of DLP client"""
    global dlp_client
    if dlp_client is None:
        try:
            dlp_client = dlp_v2.DlpServiceClient()
            logger.info("DLP client initialized successfully")
        except Exception as e:
            logger.warning(f"DLP client initialization failed: {e}")
            logger.warning("PII redaction will be disabled")
    return dlp_client

if not PROJECT_ID:
    logger.warning("GOOGLE_CLOUD_PROJECT not set; DLP may fail without ADC context.")


# Fix: Use correct DLP config objects for current google-cloud-dlp


INFO_TYPES_TO_REDACT = [
    dlp_types.InfoType(name="PERSON_NAME"),
    dlp_types.InfoType(name="EMAIL_ADDRESS"),
    dlp_types.InfoType(name="PHONE_NUMBER"),
    dlp_types.InfoType(name="STREET_ADDRESS"),
    dlp_types.InfoType(name="CREDIT_CARD_NUMBER"),
    dlp_types.InfoType(name="DATE_OF_BIRTH"),
    dlp_types.InfoType(name="US_SOCIAL_SECURITY_NUMBER"),
    # Removed INDIA_AADHAAR_ID_NUMBER - not available in all regions
]


info_type_transformations = dlp_types.InfoTypeTransformations(
    transformations=[
        {
            "info_types": INFO_TYPES_TO_REDACT,
            "primitive_transformation": dlp_types.PrimitiveTransformation(
                replace_with_info_type_config=dlp_types.ReplaceWithInfoTypeConfig()
            ),
        }
    ]
)

DEIDENTIFY_CONFIG = dlp_types.DeidentifyConfig(
    info_type_transformations=info_type_transformations
)
# --- 5B. TRANSLATION FEATURE (Firestore + Translation Support) ---

from fastapi import Query, Header
from google.cloud import firestore
from translation_utils import (
    SUPPORTED_LANGUAGES,
    translate_analysis_content,
    translate_negotiation_email
)

# Initialize Firestore client for translations
try:
    firestore_client = firestore.Client()
    logger.info("Firestore client initialized for translations")
except Exception as e:
    logger.warning(f"Firestore client initialization failed: {e}")
    firestore_client = None

# Initialize Cloud Storage client for async processing
try:
    storage_client_gcs = gcs_storage.Client()
    logger.info("Google Cloud Storage client initialized for async processing")
except Exception as e:
    logger.warning(f" Cloud Storage client initialization failed: {e}")
    storage_client_gcs = None

# Configuration for async processing
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "lexiguard-documents")
MAX_FILE_SIZE_MB = 10  # 10MB limit as per your UI

@app.get("/supported-languages")
async def get_supported_languages():
    """Return the list of supported translation languages with categories"""
    from translation_utils import get_language_categories
    
    categories = get_language_categories()
    
    # Create a simpler structure for frontend
    languages_by_category = {}
    for category, codes in categories.items():
        languages_by_category[category] = [
            {"code": code, "name": SUPPORTED_LANGUAGES.get(code, code)}
            for code in codes
            if code in SUPPORTED_LANGUAGES
        ]
    
    return {
        "total_languages": len(SUPPORTED_LANGUAGES),
        "languages": [
            {"code": code, "name": name} 
            for code, name in SUPPORTED_LANGUAGES.items()
        ],
        "categories": languages_by_category
    }

@app.get("/language-categories")
async def get_language_categories():
    """Return language categories dynamically grouped"""
    language_categories = {
        "Asian": {k: v for k, v in SUPPORTED_LANGUAGES.items() if k in ["zh", "ja", "ko"]},
        "European": {k: v for k, v in SUPPORTED_LANGUAGES.items() if k in ["fr", "de", "it", "pt", "nl", "ru"]},
        "Indian": {k: v for k, v in SUPPORTED_LANGUAGES.items() if k in ["hi", "ta", "te", "bn", "ml"]},
        "Others": {k: v for k, v in SUPPORTED_LANGUAGES.items() if k not in ["zh", "ja", "ko", "fr", "de", "it", "pt", "nl", "ru", "hi", "ta", "te", "bn", "ml"]},
    }
    return {"categories": language_categories}

# REPLACE your existing /translate/{analysis_id} endpoint with this:
@app.post("/translate/{analysis_id}")
async def translate_analysis(
    analysis_id: str,
    language: str = Query(..., description="Target language code"),
    user_id: str = Query(..., description="User ID for authorization")
):
    """
    Translate an existing analysis result (summary + risks + clauses)
    into the requested target language using Google Cloud Translation API.
    """
    logger.info(f"ðŸ”„ Translation request: {analysis_id} -> {language} (user: {user_id})")
    
    if not firestore_client:
        logger.error(" Firestore not initialized")
        raise HTTPException(status_code=500, detail="Firestore not initialized")
    
    if language not in SUPPORTED_LANGUAGES:
        logger.error(f"Unsupported language: {language}")
        available = ", ".join(list(SUPPORTED_LANGUAGES.keys())[:10])
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported language: {language}. Available: {available}..."
        )

    try:
        # Fetch analysis from Firestore
        doc_ref = firestore_client.collection("userAnalyses").document(analysis_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            logger.error(f"Analysis not found: {analysis_id}")
            raise HTTPException(status_code=404, detail="Analysis not found")

        analysis_data = doc.to_dict()
        
        # Security check - verify user owns this analysis
        if analysis_data.get("userID") != user_id:
            logger.error(f"Unauthorized access attempt by {user_id}")
            raise HTTPException(status_code=403, detail="Unauthorized access to this analysis")
        
        logger.info(f" Analysis found, checking for cached translation...")
        existing_translations = analysis_data.get("translations", {})

        # If already translated, return cached version
        if language in existing_translations:
            logger.info(f"Found cached translation for {language}")
            cached = existing_translations[language]
            
            # Ensure we return the correct structure
            response = {
                "language": language,
                "language_name": SUPPORTED_LANGUAGES.get(language, language),
                "translated_content": {
                    "summary": cached.get("summary", ""),
                    "risks": cached.get("risks", []),
                    "clauses": cached.get("clauses", []),
                    "suggestions": cached.get("suggestions", [])
                }
            }
            logger.info(f"“¦ Returning cached translation with {len(str(response))} bytes")
            return response

        logger.info(f"”„ No cache found, generating new translation for {language}")

        # Use the translate_analysis_content function from translation_utils
        translated_content = translate_analysis_content(analysis_data, language)
        
        logger.info(f" Translation completed: {translated_content.keys()}")
        
        # Validate that translation has content
        has_content = (
            translated_content.get("summary") or 
            translated_content.get("risks") or 
            translated_content.get("clauses")
        )
        
        if not has_content:
            logger.error(f"Translation returned empty content")
            logger.error(f"Translation keys: {translated_content.keys()}")
            raise HTTPException(
                status_code=500, 
                detail="Translation completed but returned empty content"
            )

        # Save translation to Firestore for caching
        try:
            existing_translations[language] = translated_content
            doc_ref.update({"translations": existing_translations})
            logger.info(f"Translation cached in Firestore for {language}")
        except Exception as e:
            logger.warning(f" Failed to cache translation: {e}")
            # Continue even if caching fails

        logger.info(f" Translation completed successfully for {language}")
        
        response = {
            "language": language,
            "language_name": SUPPORTED_LANGUAGES.get(language, language),
            "translated_content": translated_content
        }
        
        logger.info(f"“¦ Returning new translation with {len(str(response))} bytes")
        return response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f" Translation failed for {analysis_id}: {e}")
        logger.error(f" Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        import traceback
        logger.error(f" Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")







@app.get("/translation-stats/{analysis_id}")
async def translation_stats(analysis_id: str):
    """Return translation availability stats for a given analysis"""
    if not firestore_client:
        raise HTTPException(status_code=500, detail="Firestore not initialized")

    try:
        doc_ref = firestore_client.collection("userAnalyses").document(analysis_id)
        doc = doc_ref.get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="Analysis not found")

        data = doc.to_dict()
        available_languages = list(data.get("translations", {}).keys())
        remaining = max(0, len(SUPPORTED_LANGUAGES) - len(available_languages))

        return {
            "analysis_id": analysis_id,
            "available_translations": available_languages,
            "remaining_languages": remaining,
        }
    except Exception as e:
        logger.error(f"Error getting translation stats: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve stats")

# --- 6. DATA MODELS ---
class DocumentRequest(BaseModel):
    text: str

class ChatRequest(BaseModel):
    message: str
    document_text: str
    analysis_id: Optional[str] = None
    user_role: Optional[str] = None
    conversation_history: Optional[List[Dict[str, str]]] = None
    analysis_id: Optional[str] = None
    user_role: Optional[str] = None
    conversation_history: Optional[List[Dict[str, str]]] = None

class NegotiationRequest(BaseModel):
    clause: str

class DocumentEmailRequest(BaseModel):
    document_summary: str
    risk_summary: str

class SendDocumentReviewRequest(BaseModel):
    filename: str
    document_summary: str
    risk_summary: str
    clauses: list
    user_email: str

class ExtendedAnalysisRequest(BaseModel):
    text: str

# --- 7. PROMPTS (FULL RESTORED) ---

SUMMARY_PROMPT = """
You are LexiGuard, an expert AI assistant that explains complex legal documents in simple terms.

IMPORTANT: The document you're analyzing has had Personal Identifiable Information (PII) redacted and replaced with placeholders like [PERSON_NAME], [EMAIL_ADDRESS], [STREET_ADDRESS], etc. Focus on explaining the terms, obligations, and rights regardless of these placeholders.

Analyze the following contract and provide a clear, well-structured summary.

Format your response using clean markdown:
- Use ## for main section headings
- Use **bold** for key terms, amounts, and dates
- Use simple bullet points (-) for lists
- Do NOT use nested bullets or asterisks (*)
- Keep paragraphs short and readable
- Use proper spacing between sections

Structure your summary as follows:

## Document Overview
Briefly explain what type of agreement this is and identify the parties involved (use placeholders like "the Tenant" or "the Service Provider" if names are redacted).

## Key Responsibilities

**Party A Responsibilities:**
- List main obligations clearly
- Highlight any unusual or concerning terms
- Use specific numbers, dates, and amounts when mentioned

**Party B Responsibilities:**
- List main obligations clearly
- Highlight any heavy burdens or unusual requirements
- Use specific numbers, dates, and amounts when mentioned

## Duration and Important Dates
- Start date
- End date or renewal terms
- Notice periods for termination
- Any automatic renewal clauses

## Financial Terms
- Primary payment amounts and due dates
- Additional fees or charges
- Security deposits and refund conditions
- Late payment penalties
- Any provisions for fee increases

Use clear, simple language suitable for a non-lawyer. When you see redacted placeholders like [PERSON_NAME], refer to them generically (e.g., "the Tenant", "the Client", "the Service Provider").
"""

RISK_ANALYSIS_PROMPT = """
You are a meticulous risk analysis AI. Scan the provided legal document.
Your task is to identify and extract any clauses that are potentially unfavorable, non-standard, or represent a significant risk.
Focus specifically on clauses related to: Indemnity, Limitation of Liability, Automatic Renewal, Termination Penalties, and Non-Compete agreements.
For each identified clause, you MUST provide the original text, a simple one-sentence explanation of the risk, and a severity level of either 'High' or 'Medium'.
You MUST respond ONLY with a valid JSON object. The structure of the JSON object must be:
{"risks": [{"clause_text": "...", "risk_explanation": "...", "severity": "..."}]}
If you find no risks, you MUST return: {"risks": []}
Do not add any text or formatting before or after the JSON object.
"""
NEGOTIATION_PROMPT = """
You are LexiGuard, an AI that helps users politely negotiate risky contract clauses.
Draft a professional and polite email body requesting to amend or clarify the following clause.

Clause (PII may be redacted with placeholders like [PERSON_NAME]):
{clause}


Your task is to draft ONLY the email body. The email should:
1. Start professionally (e.g., "Regarding the draft agreement received...").
2. Clearly reference the clause in question (you can mention its content briefly).
3. Explain the user's concern politely based on the clause's meaning.
4. Suggest a discussion or a more balanced alternative.
5. Use standard placeholders like `[Your Name]`, `[Your Contact Info]`, `[Recipient Name/Company Name]`, and `[Date]` where specific personal or logistical details would normally go. Do NOT try to invent names or details.
6. End with a collaborative closing (e.g., "I look forward to discussing this further.").

Generate ONLY the email body text.
"""

DOCUMENT_EMAIL_PROMPT = """
You are LexiGuard, an AI assistant that helps users communicate about legal document reviews.

Generate a professional email to send to a legal advisor, counterparty, or stakeholder regarding a legal document review.

Document Summary:
{document_summary}

Identified Risks:
{risk_summary}

The email should:
1. Be professional, clear, and concise
2. Summarize the key findings from the document analysis
3. Highlight the most critical risks identified
4. Request clarification, revision, or discussion on the concerning clauses
5. Maintain a collaborative and constructive tone
6. Be ready to send with minimal editing

Generate ONLY the email body. Do not include:
- Subject line
- Sender/recipient names or addresses
- Placeholders like [Your Name] or [Company Name]

Start directly with a professional greeting and the email content.
"""

FAIRNESS_PROMPT = """
You are LexiGuard, a fairness evaluator.
Compare the following risky clause with a standard, balanced contract clause.
Return a JSON object strictly in this format:
{
  "standard_clause": "...",
  "risky_clause": "...",
  "fairness_score": 0-100,
  "explanation": "..."
}

Risky Clause:
{clause}
"""

DETAILED_CLAUSE_ANALYSIS_PROMPT = """
You are a legal expert analyzing contracts and agreements for LexiGuard. 
Analyze the following document and identify ALL risky or concerning clauses with deep explanations.

For EACH risky clause, provide:
1. **clause**: The exact clause text or relevant excerpt (keep it under 200 characters if too long)
2. **risk_level**: "High", "Medium", or "Low"
3. **impact**: Brief description of potential harm to the user (1 sentence)
4. **recommendation**: Specific actionable advice for negotiation (1-2 sentences)
5. **explanation**: Detailed plain-language explanation of why this clause is risky and what could go wrong (2-3 sentences)

Focus on these risk categories:
- Termination rights (sudden eviction, firing without cause)
- Financial liability (unlimited damages, penalties)
- Automatic renewal traps
- Non-compete restrictions
- Indemnification clauses
- Limitation of liability
- Unilateral changes by one party
- Waiver of legal rights

You MUST respond ONLY with a valid JSON array. Use this exact format:
[
    {
        "clause": "Original clause text here",
        "risk_level": "High",
        "impact": "Brief impact description",
        "recommendation": "What the user should do or negotiate",
        "explanation": "Detailed plain-language explanation of why this is risky"
    }
]

If no risks are found, return an empty array: []
CRITICAL: Respond ONLY with valid JSON, no additional text, no markdown.
"""

# --- 8. HELPER FUNCTIONS ---
# Model is already initialized above (lines 47-82) - using gemini-1.5-flash
# No need for duplicate get_working_model() function

def redact_text_with_dlp(text: str):
    if not text or not PROJECT_ID:
        logger.warning("DLP: Project ID not configured. Skipping redaction.")
        return text, False

    # Get DLP client (lazy initialization)
    client = get_dlp_client()
    if not client:
        logger.warning("DLP client not available. Skipping redaction.")
        return text, False

    parent_path = f"projects/{PROJECT_ID}/locations/global"
    item = {"value": text}

    try:
        response = client.deidentify_content(
            request={
                "parent": parent_path,
                "deidentify_config": DEIDENTIFY_CONFIG,
                "inspect_config": {"info_types": INFO_TYPES_TO_REDACT},
                "item": item,
            }
        )
        redacted = response.item.value
        changed = redacted != text
        logger.info("DLP Redaction complete.")
        return redacted, changed
    except Exception as e:
        logger.error(f"DLP failed: {e}")
        return text, False

def extract_text_from_pdf(file_stream):
    try:
        pdf = PyPDF2.PdfReader(file_stream)
        return "".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise HTTPException(status_code=500, detail="Error extracting text from PDF.")

def extract_text_from_docx(file_stream):
    try:
        doc = Document(file_stream)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=500, detail="Error extracting text from DOCX.")
    
def extract_text_from_txt(file) -> str:
    try:
        # Read text file with UTF-8 encoding
        content = file.read()
        if isinstance(content, bytes):
            text = content.decode('utf-8')
        else:
            text = content
        return text
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            file.seek(0)
            content = file.read()
            text = content.decode('latin-1')
            return text
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise HTTPException(status_code=500, detail="Error extracting text from TXT file.")
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise HTTPException(status_code=500, detail="Error extracting text from TXT file.")

# --- 9. CORE ANALYSIS LOGIC ---
def analyze_text_internal(text: str):
    if model is None:
        raise HTTPException(status_code=503, detail="AI model not initialized. Check backend logs.")
    
    try:
        redacted_text, changed = redact_text_with_dlp(text)
        prompt = f"{SUMMARY_PROMPT}\n\nDocument:\n{redacted_text}"
        response = model.generate_content(prompt)
        summary = response.text.strip()

        risk_prompt = f"{RISK_ANALYSIS_PROMPT}\n\nDocument:\n{redacted_text}"
        risk_response = model.generate_content(risk_prompt)
        try:
            # Clean JSON response
            risks_text = risk_response.text.strip().replace("```json", "").replace("```", "").strip()
            risks = json.loads(risks_text)
        except Exception as e:
            logger.error(f"Risk JSON parse error: {e}")
            risks = {"risks": []}

        return {
    "summary": summary,
    "risks": risks,
    "pii_redacted": changed,
    "redacted_text": redacted_text  # ADD THIS LINE
}
    except Exception as e:
        logger.error(f"Error in analyze_text_internal: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

def analyze_clauses_detailed_internal(text: str):
    if model is None:
        raise HTTPException(status_code=503, detail="AI model not initialized. Check backend logs.")
    
    try:
        redacted_text, changed = redact_text_with_dlp(text)
        prompt = f"{DETAILED_CLAUSE_ANALYSIS_PROMPT}\n\nDocument:\n{redacted_text}"
        response = model.generate_content(prompt)
        try:
            # Clean JSON response
            risks_text = response.text.strip().replace("```json", "").replace("```", "").strip()
            risks = json.loads(risks_text)
        except Exception as e:
            logger.error(f"Clause JSON parse error: {e}")
            risks = []
        return {
    "risks": risks, 
    "pii_redacted": changed,
    "redacted_text": redacted_text  #  ADD THIS LINE
}
    except Exception as e:
        print(f"Error in detailed clause analysis: {str(e)}")
        return {
            "risks": [],
            "pii_redacted": False,
            "redacted_text": ""
        }

def extract_text_from_pdf(file) -> str:
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# --- 8. ROUTES ---
@app.get("/")
async def root():
    return {
        "message": "LexiGuard API is running.",
        "version": "1.3.0",
        "endpoints": [
            "/analyze",
            "/analyze-file (Standard Analysis with Negotiation)",
            "/analyze-clauses (Detailed Clause Analysis)",
            "/draft-negotiation (Generate negotiation emails)",
            "/draft-document-email (Generate comprehensive document review email)",
            "/analyze-extended (Fairness scoring)",
            "/chat"
        ]
    }

@app.post("/analyze")
async def analyze_document(request: DocumentRequest):
    return analyze_text_internal(request.text)

@app.post("/analyze-file")
async def analyze_file(file: UploadFile = File(None), text: str = Form(None)):
    """
    STANDARD ANALYSIS endpoint with negotiation support.
    Returns: summary, risks, suggestions, and file type.
    """
    if file:
        filename = file.filename.lower()
        if filename.endswith(".pdf"):
            document_text = extract_text_from_pdf(file.file)
            file_type = "PDF"
        elif filename.endswith(".docx"):
            document_text = extract_text_from_docx(file.file)
            file_type = "DOCX"
        elif filename.endswith(".txt"):
            document_text = extract_text_from_txt(file.file)
            file_type = "TXT"
        else:
            return {"error": "Unsupported file type. Only PDF, DOCX, or TXT allowed."}
    elif text:
        document_text = text
        file_type = "Text"
    else:
        raise HTTPException(status_code=400, detail="No file or text provided")

    # Redact PII from document for chat (privacy protection)
    redacted_document_text, _ = redact_text_with_dlp(document_text)
    
    result = analyze_text_internal(document_text)
    
    # GENERATE DYNAMIC SUGGESTIONS using AI
    risks_list = result.get("risks", {}).get("risks", [])
    suggestions = []
    
    if risks_list and model:
        try:
            # Create a prompt to generate contextual suggestions
            risks_summary = "\n".join([
                f"- {risk.get('severity', 'Unknown')} Risk: {risk.get('risk_explanation', 'No explanation')}"
                for risk in risks_list[:5]  # Limit to top 5 risks
            ])
            
            suggestion_prompt = f"""
You are LexiGuard, an AI legal assistant. Based on the following risks identified in a legal document, provide 4-6 specific, actionable suggestions for the user.

Identified Risks:
{risks_summary}

Generate practical suggestions that:
1. Address the specific risks mentioned
2. Provide clear action items (e.g., "Request clarification on...", "Negotiate to change...", "Consult a lawyer about...")
3. Are professional and constructive
4. Help the user protect their interests

Return ONLY a JSON array of suggestion strings. Format:
["suggestion 1", "suggestion 2", "suggestion 3", ...]

Do NOT include any markdown, code blocks, or extra text - just the JSON array.
"""
            
            suggestion_response = model.generate_content(suggestion_prompt)
            suggestions_text = suggestion_response.text.strip().replace("```json", "").replace("```", "").strip()
            
            try:
                suggestions = json.loads(suggestions_text)
                if not isinstance(suggestions, list):
                    raise ValueError("Response is not a list")
            except Exception as e:
                logger.error(f"Failed to parse AI suggestions: {e}")
                # Fallback to smart default suggestions
                suggestions = generate_fallback_suggestions(risks_list)
                
        except Exception as e:
            logger.error(f"Error generating AI suggestions: {e}")
            suggestions = generate_fallback_suggestions(risks_list)
    else:
        # No risks found - provide positive suggestions
        suggestions = [
            "No significant risks detected. The document appears to contain standard terms and conditions.",
            "Proceed with standard due diligence: verify all parties' information, dates, and financial terms.",
            "Ensure you understand all obligations and responsibilities outlined in the agreement before signing.",
            "Keep a signed copy of the agreement for your records and future reference."
        ]
    
    # Return in format expected by frontend
    response = {
        "filename": file.filename if file else "Text Input",
        "file_type": file.filename.split(".")[-1].upper() if file else "Text",
        "summary": result.get("summary", ""),
        "risks": risks_list,
        "suggestions": suggestions,
        "pii_redacted": result.get("pii_redacted", False),
        "redacted_document_text": redacted_document_text
    }
    
    # Add privacy notice if PII was redacted
    if result.get("pii_redacted", False):
        response["privacy_notice"] = "“ Your Personal Data Has Been Redacted for Privacy."
    
    return response


def generate_fallback_suggestions(risks_list):
    """Generate intelligent fallback suggestions based on risk analysis"""
    suggestions = []
    
    high_risks = [r for r in risks_list if r.get("severity") == "High"]
    medium_risks = [r for r in risks_list if r.get("severity") == "Medium"]
    
    # Analyze risk types
    risk_keywords = {
        "liability": ["liability", "indemnify", "indemnification", "damages"],
        "termination": ["termination", "terminate", "cancel", "end"],
        "renewal": ["renewal", "renew", "automatic", "auto-renew"],
        "payment": ["payment", "fee", "penalty", "fine", "charge"],
        "non-compete": ["non-compete", "non compete", "restrict", "prohibition"]
    }
    
    detected_risks = set()
    for risk in risks_list:
        risk_text = (risk.get("clause_text", "") + " " + risk.get("risk_explanation", "")).lower()
        for risk_type, keywords in risk_keywords.items():
            if any(keyword in risk_text for keyword in keywords):
                detected_risks.add(risk_type)
    
    # Generate specific suggestions based on detected risk types
    if high_risks:
        suggestions.append(f"Immediate attention required: {len(high_risks)} high-risk clause(s) identified that could significantly impact your rights or obligations.")
    
    if "liability" in detected_risks:
        suggestions.append("Negotiate to add reasonable liability caps or limitations to protect yourself from unlimited financial exposure.")
    
    if "termination" in detected_risks:
        suggestions.append("Request mutual termination rights with adequate notice period to ensure fair treatment for both parties.")
    
    if "renewal" in detected_risks:
        suggestions.append("Ask to remove automatic renewal clauses or add clear opt-out procedures with sufficient advance notice.")
    
    if "payment" in detected_risks:
        suggestions.append("Clarify all payment terms, penalties, and fee structures in writing before signing the agreement.")
    
    if "non-compete" in detected_risks:
        suggestions.append("Negotiate to narrow the scope, duration, and geographic limitations of any non-compete restrictions.")
    
    # Always add general advice
    if len(risks_list) >= 3:
        suggestions.append("Strongly recommend engaging a legal professional to review all identified risks before signing this agreement.")
    
    suggestions.append("Document all communications and keep records of any proposed amendments or clarifications.")
    
    if not suggestions:
        suggestions = [
            f"Review the {len(risks_list)} identified risk(s) carefully with the other party.",
            "Consider consulting with legal counsel to address any concerns.",
            "Request written clarification for any unclear terms before proceeding."
        ]
    
    return suggestions

@app.post("/analyze-clauses")
async def analyze_clauses(file: UploadFile = File(None), text: str = Form(None)):
    """
    DETAILED CLAUSE ANALYSIS endpoint.
    Deep clause-by-clause analysis with detailed explanations,
    impact assessment, and recommendations.
    
    Returns:
    - filename (if file uploaded)
    - file_type (PDF, DOCX, TXT, or Text)
    - total_risky_clauses (count)
    - clauses (array of detailed clause objects)
    - document_preview (first 300 characters)
    """
    if file:
        filename = file.filename.lower()
        if filename.endswith(".pdf"):
            document_text = extract_text_from_pdf(file.file)
            file_type = "PDF"
        elif filename.endswith(".docx"):
            document_text = extract_text_from_docx(file.file)
            file_type = "DOCX"
        elif filename.endswith(".txt"):  # ADD TXT SUPPORT
            document_text = extract_text_from_txt(file.file)
            file_type = "TXT"
        else:
            return {"error": "Unsupported file type. Only PDF, DOCX, or TXT allowed."}
        filename_display = file.filename
    elif text:
        document_text = text
        file_type = "Text"  #  ADD file_type for manual text input
        filename_display = "Text Input"
    else:
        raise HTTPException(status_code=400, detail="No file or text provided")

    # Redact PII from document for chat (privacy protection)
    redacted_document_text, _ = redact_text_with_dlp(document_text)
    
    result = analyze_clauses_detailed_internal(document_text)
    
    # Return in format expected by frontend
    response = {
        "filename": file.filename if file else "Text Input",
        "file_type": file.filename.split(".")[-1].upper() if file else "Text",
        "total_risky_clauses": len(result.get("risks", [])),
        "clauses": result.get("risks", []),
        "pii_redacted": result.get("pii_redacted", False),
        "redacted_text": redacted_document_text
    }
    
    # Add privacy notice if PII was redacted
    if result.get("pii_redacted", False):
        response["privacy_notice"] = "“ Your Personal Data Has Been Redacted for Privacy."
    
    return response

@app.post("/negotiate-clause")
@app.post("/draft-negotiation")
async def negotiate_clause(request: NegotiationRequest):
    """Generate negotiation email for a risky clause (using redacted text)"""
    redacted_text, changed = redact_text_with_dlp(request.clause)
    prompt = NEGOTIATION_PROMPT.format(clause=redacted_text)
    response = model.generate_content(prompt)
    return {"negotiation_email": response.text.strip()}

@app.post("/generate-email")
@app.post("/draft-document-email")
async def generate_email(request: DocumentEmailRequest):
    """Generate comprehensive document review email"""
    prompt = DOCUMENT_EMAIL_PROMPT.format(
        document_summary=request.document_summary,
        risk_summary=request.risk_summary
    )
    response = model.generate_content(prompt)
    return {"document_email": response.text.strip()}

@app.post("/fairness-score")
async def fairness_score(request: NegotiationRequest):
    prompt = FAIRNESS_PROMPT.format(clause=request.clause)
    response = model.generate_content(prompt)
    try:
        return json.loads(response.text)
    except Exception:
        return {"error": "Could not parse AI response"}

# --- SEND DOCUMENT REVIEW EMAIL ENDPOINT ---
@app.post("/send-document-review")
async def send_document_review(request: SendDocumentReviewRequest):
    """
    Generate PDF document review report and send via email
    """
    try:
        # Create PDF in memory
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#064E3B'),
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph("LexiGuard Document Review Report", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Add document info
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#0891B2'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        story.append(Paragraph("Document Information", heading_style))
        story.append(Paragraph(f"<b>Filename:</b> {request.filename}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Add document summary
        story.append(Paragraph("Document Summary", heading_style))
        story.append(Paragraph(request.document_summary, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Add risk summary
        story.append(Paragraph("Risk Analysis", heading_style))
        story.append(Paragraph(request.risk_summary, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Add risky clauses table
        if request.clauses:
            story.append(Paragraph("Identified Risky Clauses", heading_style))
            
            # Create table data
            table_data = [['Clause', 'Risk Level', 'Explanation']]
            for clause_item in request.clauses:
                clause_text = clause_item.get('clause', 'N/A')[:100] + '...' if len(clause_item.get('clause', '')) > 100 else clause_item.get('clause', 'N/A')
                risk = clause_item.get('risk', 'Unknown')
                explanation = clause_item.get('explanation', 'No explanation provided')[:150] + '...' if len(clause_item.get('explanation', '')) > 150 else clause_item.get('explanation', 'No explanation provided')
                
                table_data.append([clause_text, risk, explanation])
            
            # Create table with styling
            table = Table(table_data, colWidths=[2.5*inch, 1*inch, 3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#064E3B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        pdf_data = pdf_buffer.read()
        
        # Send email using Gmail SMTP
        # Note: You'll need to set up Gmail App Password and add to .env
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = os.getenv("GMAIL_SENDER_EMAIL")  # Add to .env
        sender_password = os.getenv("GMAIL_APP_PASSWORD")  # Add Gmail App Password to .env
        
        if not sender_email or not sender_password:
            raise HTTPException(
                status_code=500,
                detail="Email configuration missing. Please set GMAIL_SENDER_EMAIL and GMAIL_APP_PASSWORD in .env file"
            )
        
        # Create email
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = request.user_email
        msg['Subject'] = f"LexiGuard Document Review: {request.filename}"
        
        # Email body
        body = f"""
Hello,

Please find attached your LexiGuard document review report for: {request.filename}

This report includes:
- Document Summary
- Risk Analysis
- Identified Risky Clauses

Thank you for using LexiGuard!

Best regards,
The LexiGuard Team
"""
        msg.attach(MIMEText(body, 'plain'))
        
        # Attach PDF
        pdf_attachment = MIMEApplication(pdf_data, _subtype='pdf')
        pdf_attachment.add_header('Content-Disposition', 'attachment', filename=f'LexiGuard_Review_{request.filename}.pdf')
        msg.attach(pdf_attachment)
        
        # Send email
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        
        return {
            "success": True,
            "message": f"Document review email sent successfully to {request.user_email}"
        }
        
    except Exception as e:
        logger.error(f"Error sending document review email: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

@app.get("/")
def root():
    return {"message": "LexiGuard API is running successfully ðŸš€"}

# --- ENHANCED CHAT ENDPOINT WITH ROLE-AWARE FUNCTIONALITY ---
@app.post("/chat")
async def chat_with_document(request: ChatRequest):
    """
    Enhanced chat endpoint supporting role-aware intelligent conversations.
    
    Features:
    - Role discovery and persistence
    - Intent routing (retrieval vs analysis)
    - Persona-based responses
    - Conversation history support
    """
    if not model:
        raise HTTPException(status_code=503, detail="AI model not initialized")
    
    logger.info(f"Chat request for analysis ID: {request.analysis_id}, Role: {request.user_role}")
    
    # Role discovery and persistence logic
    current_user_role = request.user_role
    
    # If we have an analysis_id, check for stored role in Firestore
    if request.analysis_id and firestore_client:
        try:
            doc_ref = firestore_client.collection("userAnalyses").document(request.analysis_id)
            doc = doc_ref.get()
            if doc.exists:
                stored_role = doc.to_dict().get("userRole")
                if stored_role and not current_user_role:
                    current_user_role = stored_role
                    logger.info(f"Retrieved stored role: {stored_role}")
        except Exception as e:
            logger.warning(f"Error accessing Firestore for role: {e}")
    
    # Role discovery prompts
    INITIAL_ROLE_PROMPT = (
        "Hello! I'm here to help you understand this document. "
        "To give you the most relevant and personalized insights, "
        "could you please tell me **your role** in this document? "
        "(e.g., Tenant, Landlord, Employee, Employer, Borrower, Lender, Seller, Buyer, Freelancer, Client, Party A, etc.)"
    )
    
    ROLE_ACKNOWLEDGEMENT_PROMPT = (
        "Understood. I will answer your questions from the perspective of the **{user_role}** in this document. "
        "Now, what's your first question about this document?"
    )
    
    # Check if we need to discover the user's role
    if not current_user_role:
        # Check if user's message might be a role declaration
        possible_role_input = request.message.strip()
        role_keywords = ["tenant", "employee", "borrower", "landlord", "employer", "lender", 
                        "party", "seller", "buyer", "client", "freelancer", "contractor", 
                        "service provider", "customer", "user", "member"]
        
        is_likely_role_answer = (
            len(possible_role_input.split()) < 5 and
            any(kw in possible_role_input.lower() for kw in role_keywords)
        )
        
        if is_likely_role_answer:
            # User is declaring their role
            user_role_declared = possible_role_input
            
            # Save this role to Firestore if analysis_id is provided
            if request.analysis_id and firestore_client:
                try:
                    doc_ref = firestore_client.collection("userAnalyses").document(request.analysis_id)
                    doc_ref.update({"userRole": user_role_declared})
                    logger.info(f"Role '{user_role_declared}' saved for analysis ID {request.analysis_id}")
                except Exception as e:
                    logger.warning(f"Failed to save role to Firestore: {e}")
            
            return {
                "reply": ROLE_ACKNOWLEDGEMENT_PROMPT.format(user_role=user_role_declared),
                "identified_role": user_role_declared,
                "needs_role_input": False
            }
        else:
            # Role is unknown, ask for it
            return {
                "reply": INITIAL_ROLE_PROMPT,
                "identified_role": None,
                "needs_role_input": True
            }
    
    # Intent routing using Gemini
    router_prompt = f"""
    You are an AI assistant routing user queries about a legal document.
    The user's role is '{current_user_role}'.
    Classify the user's question into one of the following categories, responding ONLY with the category name:
    1. Retrieval: Question asks for a specific fact, definition, or clause from the document.
    2. Analysis: Question asks for explanation, opinion, implication, risk assessment, or clarification related to the document (from the user's role perspective).
    3. General: Question is outside the scope of the document (e.g., general knowledge, small talk).

    User Question: "{request.message}"
    """
    
    try:
        router_response = model.generate_content(router_prompt)
        intent = router_response.text.strip().lower()
        logger.info(f"Router classified intent: {intent}")
    except Exception as e:
        logger.error(f"Router error: {e}")
        intent = "general"
    
    # Persona-based system prompts
    SYSTEM_PROMPT_RETRIEVAL = f"""
    You are a highly accurate, factual assistant providing answers strictly based on the provided document.
    The user is interacting as the **'{current_user_role}'**. When asked a factual question, extract the exact relevant information from the document.
    If the information is not explicitly in the document, state clearly that you cannot find it in the provided text.
    Do NOT offer opinions or external advice.
    """
    
    SYSTEM_PROMPT_ANALYSIS = f"""
    You are LexiGuard, an expert AI legal co-pilot. Your primary goal is to help the user understand and navigate this legal document.
    The user identifies as the **'{current_user_role}'**. Always provide insights and explanations *from this perspective*.
    Explain complex clauses in simple, easy-to-understand language. Identify potential implications and suggest considerations relevant to the user's role.
    You are supportive and empowering, but always clarify that you are an AI and cannot provide legal advice.
    """
    
    SYSTEM_PROMPT_GENERAL = """
    You are LexiGuard, a helpful AI assistant. The user has asked a general question that may not be directly related to their document.
    Provide a helpful response while gently steering the conversation back to document-related topics if appropriate.
    """
    
    # Build conversation context for Gemini
    conversation_context = f"""
Document Context (PII may be redacted):
{request.document_text}

User Role: {current_user_role}
"""
    
    # Add conversation history if provided
    if request.conversation_history:
        conversation_context += "\n\nPrevious Conversation:\n"
        for msg in request.conversation_history[-4:]:  # Last 4 exchanges
            sender = "User" if msg.get("sender") == "user" else "Assistant"
            conversation_context += f"{sender}: {msg.get('text', '')}\n"
    
    conversation_context += f"\n\nCurrent User Question: {request.message}"
    
    # Generate response based on intent
    response_text = "I'm sorry, I'm having trouble processing that right now. Please try again."
    
    try:
        if "retrieval" in intent:
            prompt = f"{SYSTEM_PROMPT_RETRIEVAL}\n\n{conversation_context}"
            response = model.generate_content(prompt)
            response_text = response.text
            
        elif "analysis" in intent:
            prompt = f"{SYSTEM_PROMPT_ANALYSIS}\n\n{conversation_context}"
            response = model.generate_content(prompt)
            response_text = response.text
            
        else:  # General or fallback
            prompt = f"{SYSTEM_PROMPT_GENERAL}\n\n{conversation_context}"
            response = model.generate_content(prompt)
            response_text = response.text
            
    except Exception as e:
        logger.error(f"Error generating response for intent '{intent}': {e}")
        response_text = "I'm sorry, I'm currently unable to provide that information. Please try again later."
    
    return {
        "reply": response_text,
        "identified_role": current_user_role,
        "needs_role_input": False,
        "intent": intent
    }

@app.post("/analyze-file-async")
async def analyze_file_async(
    file: UploadFile = File(...),
    documentTitle: str = Form(...),
    analysisType: str = Form("standard"),
    userId: str = Form(...)  # Get from Firebase Auth on frontend
):
    """
    ðŸš€ NEW ASYNC ENDPOINT for queued processing
    
    Uploads file to Cloud Storage and creates a job for background processing.
    Returns immediately with job ID (non-blocking).
    
    Frontend should poll /job-status/{jobId} or use Firestore real-time listener.
    """
    try:
        logger.info(f"ðŸ“¤ Async upload request from user: {userId}")
        logger.info(f"   Document: {documentTitle}")
        logger.info(f"   Analysis: {analysisType}")
        
        # Validate Cloud Storage client
        if not storage_client_gcs:
            raise HTTPException(
                status_code=503,
                detail="Cloud Storage not configured. Please set GCS_BUCKET_NAME in environment."
            )
        
        if not firestore_client:
            raise HTTPException(
                status_code=503,
                detail="Firestore not configured for async processing."
            )
        
        # Read file content
        file_content = await file.read()
        file_size_mb = len(file_content) / (1024 * 1024)
        
        # Validate file size
        if file_size_mb > MAX_FILE_SIZE_MB:
            raise HTTPException(
                status_code=400,
                detail=f"File size ({file_size_mb:.2f}MB) exceeds {MAX_FILE_SIZE_MB}MB limit"
            )
        
        # Validate file type
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.pdf') or 
                filename_lower.endswith('.docx') or 
                filename_lower.endswith('.txt')):
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Only PDF, DOCX, and TXT are allowed."
            )
        
        # Determine file type
        if filename_lower.endswith('.pdf'):
            file_type = "pdf"
        elif filename_lower.endswith('.docx'):
            file_type = "docx"
        else:
            file_type = "txt"
        
        # Generate unique job ID
        job_id = str(uuid.uuid4())
        
        # Generate GCS path: uploads/{userId}/{timestamp}_{filename}
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        safe_filename = file.filename.replace(' ', '_')  # Remove spaces
        gcs_path = f"uploads/{userId}/{timestamp}_{safe_filename}"
        
        logger.info(f"ðŸ“¤ Uploading to GCS: gs://{GCS_BUCKET_NAME}/{gcs_path}")
        
        # Upload file to Cloud Storage
        try:
            bucket = storage_client_gcs.bucket(GCS_BUCKET_NAME)
            blob = bucket.blob(gcs_path)
            
            # Upload with metadata
            blob.upload_from_string(
                file_content,
                content_type=file.content_type
            )
            
            logger.info(f" File uploaded to Cloud Storage successfully")
            
        except Exception as upload_error:
            logger.error(f" GCS upload failed: {upload_error}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to upload file to Cloud Storage: {str(upload_error)}"
            )
        
        # Create job entry in Firestore (analysisJobs collection)
        # This will trigger the Cloud Function to publish to Pub/Sub
        try:
            job_data = {
                'jobId': job_id,
                'userID': userId,
                'documentTitle': documentTitle,
                'originalFilename': file.filename,
                'fileType': file_type,
                'gcsPath': gcs_path,
                'status': 'pending',  # Will trigger Cloud Function
                'analysisType': analysisType,
                'createdAt': firestore.SERVER_TIMESTAMP,
                'updatedAt': firestore.SERVER_TIMESTAMP,
            }
            
            # Save to Firestore
            job_ref = firestore_client.collection('analysisJobs').document(job_id)
            job_ref.set(job_data)
            
            logger.info(f" Job created in Firestore: {job_id}")
            
        except Exception as firestore_error:
            logger.error(f" Firestore job creation failed: {firestore_error}")
            
            # Clean up uploaded file
            try:
                blob.delete()
                logger.info("ðŸ§¹ Cleaned up uploaded file after Firestore error")
            except:
                pass
            
            raise HTTPException(
                status_code=500,
                detail=f"Failed to create analysis job: {str(firestore_error)}"
            )
        
        # Return job ID immediately (non-blocking)
        return {
            "success": True,
            "message": "File uploaded successfully. Analysis in progress...",
            "jobId": job_id,
            "status": "pending",
            "estimatedTime": "30-60 seconds",
            "documentTitle": documentTitle,
            "fileType": file_type
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f" Async upload error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Upload failed: {str(e)}"
        )


@app.get("/job-status/{job_id}")
async def get_job_status(job_id: str, user_id: str = Query(...)):
    """
    Get the status of an async analysis job
    
    Status values:
    - pending: Job created, waiting for worker
    - processing: Worker is processing the document
    - completed: Analysis complete, results available
    - failed: Processing failed with error
    """
    try:
        if not firestore_client:
            raise HTTPException(status_code=503, detail="Firestore not configured")
        
        # Get job from Firestore
        job_ref = firestore_client.collection('analysisJobs').document(job_id)
        job_doc = job_ref.get()
        
        if not job_doc.exists:
            raise HTTPException(status_code=404, detail="Job not found")
        
        job_data = job_doc.to_dict()
        
        # Security check - verify user owns this job
        if job_data.get('userID') != user_id:
            raise HTTPException(status_code=403, detail="Unauthorized access to this job")
        
        response = {
            "jobId": job_id,
            "status": job_data.get('status', 'unknown'),
            "documentTitle": job_data.get('documentTitle', ''),
            "createdAt": job_data.get('createdAt'),
            "updatedAt": job_data.get('updatedAt'),
        }
        
        # Add result data if completed
        if job_data.get('status') == 'completed':
            response['resultAnalysisId'] = job_data.get('resultAnalysisId')
            response['processingTimeSeconds'] = job_data.get('processingTimeSeconds')
        
        # Add error message if failed
        if job_data.get('status') == 'failed':
            response['errorMessage'] = job_data.get('errorMessage', 'Unknown error')
        
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f" Error fetching job status: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch job status: {str(e)}")


@app.get("/analysis-result/{analysis_id}")
async def get_analysis_result(analysis_id: str, user_id: str = Query(...)):
    """
    Get the full analysis results from userAnalyses collection
    
    Called after job status shows 'completed'
    """
    try:
        if not firestore_client:
            raise HTTPException(status_code=503, detail="Firestore not configured")
        
        # Get analysis from Firestore
        analysis_ref = firestore_client.collection('userAnalyses').document(analysis_id)
        analysis_doc = analysis_ref.get()
        
        if not analysis_doc.exists:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        analysis_data = analysis_doc.to_dict()
        
        # Security check - verify user owns this analysis
        if analysis_data.get('userID') != user_id:
            raise HTTPException(status_code=403, detail="Unauthorized access to this analysis")
        
        # Return in format compatible with your frontend
        return {
            "filename": analysis_data.get('originalFilename', ''),
            "file_type": analysis_data.get('fileType', '').upper(),
            "summary": analysis_data.get('summary', ''),
            "risks": analysis_data.get('risks', []),
            "recommendations": analysis_data.get('recommendations', []),
            "clauseAnalysis": analysis_data.get('clauseAnalysis', {}),
            "pii_redacted": analysis_data.get('piiRedacted', False),
            "redacted_document_text": analysis_data.get('redactedDocumentText', ''),
            "analysisType": analysis_data.get('analysisType', 'standard'),
            "uploadTimestamp": analysis_data.get('uploadTimestamp'),
            "processingTimeSeconds": analysis_data.get('processingTimeSeconds')
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f" Error fetching analysis result: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch analysis: {str(e)}")


# --- TEST ENDPOINT FOR GEMINI API ---
@app.get("/test-gemini")
async def test_gemini():
    """Test endpoint to verify Gemini API is working correctly."""
    if not model:
        return {
            "status": "error",
            "message": "Gemini model not initialized",
            "api_key_configured": bool(API_KEY),
            "model_name": MODEL_NAME,
            "suggestion": "Check server logs for initialization errors"
        }
    
    try:
        # Simple test prompt
        test_response = model.generate_content("Hello! Please respond with 'Gemini API is working correctly.'")
        return {
            "status": "success",
            "message": "Gemini API is working correctly",
            "api_key_configured": True,
            "model_name": MODEL_NAME,
            "test_response": test_response.text[:100]
        }
    except Exception as e:
        error_msg = str(e)
        return {
            "status": "error",
            "message": f"Gemini API test failed: {error_msg}",
            "api_key_configured": bool(API_KEY),
            "model_name": MODEL_NAME,
            "error_type": "authentication" if "403" in error_msg else "model_access" if "404" in error_msg else "unknown"
        }


# --- RUN SERVER ---
if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))  #  Changed from 8000 to 8080
    logger.info(f"Starting server on host=0.0.0.0, port={port}")
    uvicorn.run("main:app", host="0.0.0.0", port=port, log_level="info")